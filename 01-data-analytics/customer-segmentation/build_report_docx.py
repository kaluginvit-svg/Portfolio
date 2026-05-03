# -*- coding: utf-8 -*-
"""
Скрипт для сборки отчёта в формате DOCX по сегментации User_Data_Dataset.
Генерирует графики и формирует документ с отчётом, кодом и иллюстрациями.
"""
import os
import json
import sys

# Рабочая директория — папка с данными
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
os.chdir(SCRIPT_DIR)

def load_notebook():
    with open("User_Segmentation_Analysis.ipynb", "r", encoding="utf-8") as f:
        return json.load(f)

def extract_code_and_report(nb):
    code_parts = []
    report_lines = []
    for cell in nb["cells"]:
        if cell["cell_type"] == "code":
            src = cell.get("source", [])
            code_parts.append("".join(src) if isinstance(src, list) else src)
        if cell["cell_type"] == "markdown":
            src = cell.get("source", [])
            text = "".join(src) if isinstance(src, list) else src
            if "Краткий отчёт" in text or "Краткий отчет" in text:
                report_lines.append(text)
    full_code = "\n\n".join(code_parts)
    report_text = report_lines[0] if report_lines else ""
    return full_code, report_text

def run_analysis_and_plot():
    import pandas as pd
    import numpy as np
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    df = pd.read_csv("User_Data_Dataset.csv")
    numeric_cols = ["TimeOnSite", "PagesViewed", "ItemsInCart", "Purchases"]
    for col in numeric_cols:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].median())

    def add_conversion(df_grouped, purchases_col="Purchases", cart_col="ItemsInCart"):
        df_grouped = df_grouped.copy()
        df_grouped["Conversion"] = df_grouped.apply(
            lambda row: row[purchases_col] / row[cart_col] if row[cart_col] > 0 else 0,
            axis=1,
        )
        return df_grouped

    geo_country_stats = df.groupby("Country").agg(
        TimeOnSite_mean=("TimeOnSite", "mean"),
        PagesViewed_mean=("PagesViewed", "mean"),
        Purchases_sum=("Purchases", "sum"),
        ItemsInCart_sum=("ItemsInCart", "sum"),
    ).reset_index()
    geo_country_stats = add_conversion(
        geo_country_stats, purchases_col="Purchases_sum", cart_col="ItemsInCart_sum"
    )

    device_stats = df.groupby("DeviceType").agg(
        TimeOnSite_mean=("TimeOnSite", "mean"),
        PagesViewed_mean=("PagesViewed", "mean"),
        Purchases_sum=("Purchases", "sum"),
        ItemsInCart_sum=("ItemsInCart", "sum"),
    ).reset_index()
    device_stats = add_conversion(
        device_stats, purchases_col="Purchases_sum", cart_col="ItemsInCart_sum"
    )

    os_stats = df.groupby("OS").agg(
        TimeOnSite_mean=("TimeOnSite", "mean"),
        PagesViewed_mean=("PagesViewed", "mean"),
        Purchases_sum=("Purchases", "sum"),
        ItemsInCart_sum=("ItemsInCart", "sum"),
    ).reset_index()
    os_stats = add_conversion(
        os_stats, purchases_col="Purchases_sum", cart_col="ItemsInCart_sum"
    )

    traffic_stats = df.groupby("TrafficSource").agg(
        TimeOnSite_mean=("TimeOnSite", "mean"),
        PagesViewed_mean=("PagesViewed", "mean"),
        Purchases_sum=("Purchases", "sum"),
        ItemsInCart_sum=("ItemsInCart", "sum"),
    ).reset_index()
    traffic_stats = add_conversion(
        traffic_stats, purchases_col="Purchases_sum", cart_col="ItemsInCart_sum"
    )

    fig_dir = "report_figures"
    os.makedirs(fig_dir, exist_ok=True)

    # 1. Круговая — пользователи по странам
    users_by_country = df.groupby("Country").size().reset_index(name="Users")
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.pie(
        users_by_country["Users"],
        labels=users_by_country["Country"],
        autopct="%1.1f%%",
        startangle=90,
    )
    ax.set_title("Распределение пользователей по странам")
    fig.savefig(os.path.join(fig_dir, "fig1_pie_country.png"), dpi=150, bbox_inches="tight")
    plt.close(fig)

    # 2. Топ-5 стран — три столбчатых
    top5_countries = df.groupby("Country").size().nlargest(5).index.tolist()
    top5 = geo_country_stats[geo_country_stats["Country"].isin(top5_countries)]
    top5 = top5.set_index("Country").loc[top5_countries].reset_index()

    fig, axes = plt.subplots(1, 3, figsize=(14, 4))
    axes[0].bar(top5["Country"], top5["TimeOnSite_mean"])
    axes[0].set_title("Среднее время на сайте (топ-5 стран)")
    axes[0].set_ylabel("TimeOnSite (среднее)")
    axes[0].tick_params(axis="x", rotation=45)
    axes[1].bar(top5["Country"], top5["PagesViewed_mean"])
    axes[1].set_title("Среднее кол-во страниц (топ-5 стран)")
    axes[1].set_ylabel("PagesViewed (среднее)")
    axes[1].tick_params(axis="x", rotation=45)
    axes[2].bar(top5["Country"], top5["Conversion"])
    axes[2].set_title("Конверсия (топ-5 стран)")
    axes[2].set_ylabel("Conversion")
    axes[2].tick_params(axis="x", rotation=45)
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "fig2_bar_top5.png"), dpi=150, bbox_inches="tight")
    plt.close(fig)

    # 3. Конверсия по типу устройства
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(device_stats["DeviceType"], device_stats["Conversion"])
    ax.set_title("Конверсия по типу устройства")
    ax.set_xlabel("DeviceType")
    ax.set_ylabel("Conversion")
    for bar in bars:
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{bar.get_height():.2f}",
            ha="center",
            va="bottom",
        )
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "fig3_bar_device.png"), dpi=150, bbox_inches="tight")
    plt.close(fig)

    # 4. Конверсия по ОС
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(os_stats["OS"], os_stats["Conversion"])
    ax.set_title("Конверсия по операционной системе")
    ax.set_xlabel("OS")
    ax.set_ylabel("Conversion")
    for bar in bars:
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{bar.get_height():.2f}",
            ha="center",
            va="bottom",
        )
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "fig4_bar_os.png"), dpi=150, bbox_inches="tight")
    plt.close(fig)

    # 5. Круговая — источники трафика
    users_by_traffic = df.groupby("TrafficSource").size().reset_index(name="Users")
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.pie(
        users_by_traffic["Users"],
        labels=users_by_traffic["TrafficSource"],
        autopct="%1.1f%%",
        startangle=90,
    )
    ax.set_title("Распределение пользователей по источникам трафика")
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "fig5_pie_traffic.png"), dpi=150, bbox_inches="tight")
    plt.close(fig)

    # 6. Конверсия по источникам трафика
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(traffic_stats["TrafficSource"], traffic_stats["Conversion"])
    ax.set_title("Конверсия по источникам трафика")
    ax.set_xlabel("TrafficSource")
    ax.set_ylabel("Conversion")
    for bar in bars:
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{bar.get_height():.2f}",
            ha="center",
            va="bottom",
        )
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "fig6_bar_traffic.png"), dpi=150, bbox_inches="tight")
    plt.close(fig)

    return os.path.join(SCRIPT_DIR, fig_dir)

def build_docx(code_text, report_text, fig_dir):
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit("Установите python-docx: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Заголовок
    title = doc.add_heading("Отчёт по сегментации пользователей User_Data_Dataset", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "В отчёте представлены процесс сегментации пользовательских данных из файла User_Data_Dataset.csv, "
        "расчёт ключевых метрик по сегментам (география, тип устройства, ОС, источник трафика), "
        "код на Python и графики, иллюстрирующие результаты."
    )

    # 1. Краткий отчёт
    doc.add_heading("1. Краткий отчёт", level=1)
    for block in report_text.replace("\r\n", "\n").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("##"):
            doc.add_heading(block.lstrip("#").strip(), level=2)
            continue
        # Обработка блока построчно (для смешанного **Выводы:** и списка - ...)
        lines = block.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("**") and ":**" in line:
                parts = line.replace("**", "").split(":", 1)
                p = doc.add_paragraph()
                run_bold = p.add_run(parts[0].strip() + ": ")
                run_bold.bold = True
                if len(parts) > 1:
                    p.add_run(parts[1].strip())
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    # 2. Код Python
    doc.add_heading("2. Код на Python", level=1)
    doc.add_paragraph(
        "Ниже приведён код, использованный для загрузки данных, расчёта метрик сегментации и построения графиков."
    )
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Cm(0.5)
    p_code.paragraph_format.space_before = Pt(6)
    run = p_code.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # 3. Графики и диаграммы
    doc.add_heading("3. Графики и диаграммы результатов сегментации", level=1)
    captions = [
        ("fig1_pie_country.png", "Рис. 1. Распределение пользователей по странам (круговая диаграмма)."),
        ("fig2_bar_top5.png", "Рис. 2. Топ-5 стран: среднее время на сайте, среднее число страниц, конверсия."),
        ("fig3_bar_device.png", "Рис. 3. Конверсия по типу устройства."),
        ("fig4_bar_os.png", "Рис. 4. Конверсия по операционной системе."),
        ("fig5_pie_traffic.png", "Рис. 5. Распределение пользователей по источникам трафика."),
        ("fig6_bar_traffic.png", "Рис. 6. Конверсия по источникам трафика."),
    ]
    for fname, caption in captions:
        path = os.path.join(fig_dir, fname)
        if os.path.isfile(path):
            doc.add_paragraph(caption)
            doc.add_picture(path, width=Cm(14))
            doc.add_paragraph()

    out_path = os.path.join(SCRIPT_DIR, "Отчёт_сегментация_User_Data_Dataset.docx")
    doc.save(out_path)
    return out_path

def main():
    nb = load_notebook()
    code_text, report_text = extract_code_and_report(nb)
    if not report_text:
        report_text = (
            "## 6. Краткий отчёт по сегментации User_Data_Dataset\n\n"
            "**Признаки сегментации:** Country, City (Country–City), DeviceType, OS, TrafficSource.\n\n"
            "**Метрики:** среднее время на сайте (mean TimeOnSite), среднее количество просмотренных страниц (mean PagesViewed), "
            "сумма покупок (sum Purchases), сумма добавлений в корзину (sum ItemsInCart), "
            "конверсия (Conversion = Purchases / ItemsInCart при ItemsInCart > 0).\n\n"
            "**Выводы:**\n"
            "- По географии — страны и города с наибольшей конверсией и вовлеченностью можно считать приоритетными для рекламы и контента.\n"
            "- По устройствам — тип устройства влияет на конверсию; обычно десктоп/ноутбук дают более высокую конверсию.\n"
            "- По ОС — сегменты Windows, macOS, iOS, Android показывают разную конверсию; наиболее перспективные ОС стоит учитывать при таргетинге.\n"
            "- По источникам трафика — Direct и поисковики часто дают более высокую конверсию; соцсети — выше вовлеченность. "
            "Наиболее перспективными считаются сегменты с высокой конверсией и стабильной вовлеченностью."
        )
    fig_dir = run_analysis_and_plot()
    out_path = build_docx(code_text, report_text, fig_dir)
    print("Готово. Отчёт сохранён:", out_path)

if __name__ == "__main__":
    main()
